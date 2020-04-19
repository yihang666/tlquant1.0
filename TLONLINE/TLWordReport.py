from docx import Document
from docx.shared import Inches
from numpy import unicode
import tushare as ts
from jqdatasdk import *
auth("15336553896","553896")
import pandas as pd
import time
import matplotlib.pyplot as plt
import datetime
import TLAPI
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm


#日期与实践模块

nowTime = datetime.datetime.now()
beforeTime = nowTime - datetime.timedelta(days=29)
beforeHalfYearTime = nowTime -datetime.timedelta(days=180)
yesterdayTime =  nowTime - datetime.timedelta(days=1)
#给tushare喂的数据
nowDate = nowTime.date().strftime("%Y%m%d")
beforeDate = beforeTime.date().strftime("%Y%m%d")
beforeHalfYearDate = beforeHalfYearTime.date().strftime("%Y%m%d")
yesterdayDate = yesterdayTime.date().strftime("%Y%m%d")
#给聚宽喂的数据
nowDateJQ = nowTime.date().strftime("%Y-%m-%d")
beforeDateJQ = beforeTime.date().strftime("%Y-%m-%d")
beforeHalfYearDateJQ = beforeHalfYearTime.date().strftime("%Y-%m-%d")
yesterdayDateJQ = yesterdayTime.date().strftime("%Y-%m-%d")

token = '624f1a291a6d3bec690c866a2ef0c4cc3876e11743ae207221bfed19'
ts.set_token(token)


pro = ts.pro_api()


def getDataAndSavePic():
    northMoney = TLAPI.NorthMoney()
    northmoney_file_path = r'F:\temp\northmoney.png'  # 北向资金图片地址
    plt.xticks(rotation=45)
    plt.savefig(northmoney_file_path)

    southMoney = TLAPI.SouthMoney()
    southMoney_file_path = r'F:\temp\southmoney.png'  # 南向资金绘图
    plt.savefig(southMoney_file_path)

    shibor = TLAPI.Shibor()
    shibor_file_path = r'F:\temp\shibor.png'
    plt.xticks(rotation=45)
    plt.savefig(shibor_file_path)

    current = TLAPI.Current()
    current_file_path = r'F:\temp\currency.png'
    plt.xticks(rotation=45)
    plt.savefig(current_file_path)

    wenzhouindex = TLAPI.WenzhouIndex()
    wenzhouindex_file_path = r'F:\temp\wenzhouindex.png'
    plt.savefig(wenzhouindex_file_path)

    industryVolumeRank = TLAPI.IndustryVolumeRank()
    industryVolumeRank_file_path = r'F:\temp\industryVolumeRank.png'
    plt.xticks(rotation=45)
    plt.savefig(industryVolumeRank_file_path, dpi=200, bbox_inches='tight')

    industryVolumeWeekRank = TLAPI.IndustryVolumeWeekRank()
    industryVolumeWeekRank_file_path = r'F:\temp\industryVolumeWeekRank.png'
    plt.xticks(rotation=45)
    plt.savefig(industryVolumeWeekRank_file_path, dpi=200, bbox_inches='tight')

    industryVolumeWeekRankL2 = TLAPI.IndustryVolumeWeekRankSW_L2()
    industryVolumeWeekRankL2_file_path = r'F:\temp\industryVolumeWeekRankL2.png'
    plt.xticks(rotation=45)
    plt.savefig(industryVolumeWeekRankL2_file_path, dpi=200, bbox_inches='tight')


    margin = TLAPI.Margin()
    margin_file_path = r"F:\temp\margin.png"
    plt.xticks(rotation=45)
    plt.savefig(margin_file_path,dpi=200, bbox_inches='tight')

    huanshou = TLAPI.FullChangeHand()
    huanshou_file_path = r"F:\temp\huannshou.png"
    plt.xticks(rotation=45)
    plt.savefig(huanshou_file_path, dpi=200, bbox_inches='tight')

    return (northmoney_file_path, southMoney_file_path, shibor_file_path, current_file_path, wenzhouindex_file_path,industryVolumeRank_file_path,industryVolumeWeekRank_file_path,
            margin_file_path,huanshou_file_path,industryVolumeWeekRankL2_file_path)


# %%
doc_file_path = r'F:\temp\TLDailyReport.docx'


def gen_docfile(northmoney_flie_path, southmoney_flie_path, shibor_file_path, current_file_path, wenzhouindex_file_path,industryVolumeRank_file_path,industryVolumeWeekRank_file_path,
                margin_file_path,huanshou_file_path,industryVolumeWeekRankL2_file_path,doc_file_path):
    '''
    :param df_result: 数据记录，用于表格显示
    :param pie_file_path: 饼图文件显示
    :param bar_file_path: 柱状图文件显示
    :param doc_file_path: 需要保存的WORK文件路径
    :return: 无返回值
    '''
    # 新建一个文档
    document = Document()
    document.add_heading(u' 套牢每日观察 ', 0)
    # 添加一个段落
    p = document.add_paragraph(u'作者：涨不动')
    date = str(nowDateJQ)
    document.add_paragraph(u'日期：%s' % date)
    document.add_paragraph(u'每日观察版本：1.0')

    #document.add_heading(u'重点关注涨跌幅度')

    # 今日数据复盘
    document.add_heading(u'今日数据复盘', level=1)
    document.add_heading(u'北向资金', level=2)
    document.add_picture(northmoney_flie_path, width=Inches(6.0))

    #document.add_heading(u'南向资金', level=2)
    #document.add_picture(southmoney_flie_path, width=Inches(6.0))

    document.add_heading(u'上海银行间同业拆放利率', level=2)
    document.add_picture(shibor_file_path)

    document.add_heading(u'美元汇率', level=2)
    document.add_picture(current_file_path,width=Inches(6.0))

    #document.add_heading(u'民间借贷利率(温州)', level=2)
    #document.add_picture(wenzhouindex_file_path,width=Inches(6.0))

    #document.add_heading(u'概念板块放量', level=2)
    #conceptVolumeRank = TLAPI.ConceptVolumeRank()
    # need

    document.add_heading(u'每日行业板块放量(申万一级)', level=2)
    document.add_picture(industryVolumeRank_file_path,width=Inches(6.0))

    document.add_heading(u'每周（动态）行业板块放量(申万一级)', level=2)
    document.add_picture(industryVolumeWeekRank_file_path, width=Inches(6.0))

    document.add_heading(u'每周（动态）行业板块放量(申万二级)', level=2)
    document.add_picture(industryVolumeWeekRankL2_file_path, width=Inches(6.0))

    document.add_heading(u'融资余额变动情况',level = 2)
    document.add_picture(margin_file_path,width=Inches(6.0))

    document.add_heading(u'重要指数换手率',level=2)
    document.add_picture(huanshou_file_path,width=Inches(6.0))
    '''
    tb = document.add_table(rows=len(dfHuanShou.index) + 1, cols=len(dfHuanShou.columns))
    tb.add_row()
    for i in range(len(dfHuanShou.columns)):
        tb.cell(0, i).text = dfHuanShou.columns[i]  # 添加表头
    for row in range(1, len(dfHuanShou.index)):
        for col in range(len(dfHuanShou.columns)):
            tb.cell(row, col).text = dfHuanShou.iloc[row, col]
    tb.style = 'Medium Grid 1 Accent 1'
    tb.autofit = True
    '''

    #document.add_heading(u'本月重要数据', level=1)

    document.add_heading(u'当日中国金融数据事件', level=1)
    dfTodayEvent = pro.eco_cal(date = nowDate, country = '中国')
    tb = document.add_table(rows=len(dfTodayEvent.index) + 1, cols=len(dfTodayEvent.columns))
    tb.add_row()
    for i in range(len(dfTodayEvent.columns)):
        tb.cell(0, i).text = dfTodayEvent.columns[i]  # 添加表头
    for row in range(1, len(dfTodayEvent.index)):
        for col in range(len(dfTodayEvent.columns)):
            # tb.cell(row, col).width = 1
            tb.cell(row, col).text = dfTodayEvent.iloc[row, col]
            # tb.cell(row, col).width = Cm(6)
            # tb.cell(row,col).ipynb
    tb.style = 'Medium Grid 1 Accent 1'
    tb.autofit = True


    document.add_page_break()
    document.save(doc_file_path)


(northmoney_flie_path, southmoney_flie_path, shibor_file_path, current_file_path,
 wenzhouindex_file_path,industryVolumeRank_file_path,industryVolumeWeekRank_file_path,margin_file_path,huanshou_file_path,industryVolumeWeekRankL2_file_path) = getDataAndSavePic()


gen_docfile(northmoney_flie_path, southmoney_flie_path, shibor_file_path, current_file_path, wenzhouindex_file_path,industryVolumeRank_file_path,
            industryVolumeWeekRank_file_path,margin_file_path,huanshou_file_path,industryVolumeWeekRankL2_file_path,doc_file_path)
